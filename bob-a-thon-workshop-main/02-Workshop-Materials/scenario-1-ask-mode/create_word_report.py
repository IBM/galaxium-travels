"""
Script to create a Microsoft Word document from the security analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_security_report():
    """Create comprehensive security report in Word format"""
    
    # Create document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "Security Vulnerability Analysis & Proof of Concept Report"
    doc.core_properties.author = "Bob (AI Security Analyst)"
    doc.core_properties.subject = "Legacy Inventory Management System Security Assessment"
    
    # Title Page
    title = doc.add_heading('Security Vulnerability Analysis & Proof of Concept Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Legacy Inventory Management System')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Document metadata
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    
    meta_data = [
        ('Document Classification:', 'CONFIDENTIAL - Internal Security Assessment'),
        ('Date:', 'February 5, 2026'),
        ('Prepared By:', 'Bob (AI Security Analyst)'),
        ('System Under Review:', 'Legacy Inventory Management System'),
        ('Overall Risk Rating:', '🔴 CRITICAL')
    ]
    
    for i, (key, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[1].text = value
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    
    doc.add_heading('Overview', 2)
    doc.add_paragraph(
        'This report documents critical security vulnerabilities discovered in the Legacy Inventory '
        'Management System\'s db_config.py and app.py files. The assessment reveals multiple '
        'critical-severity vulnerabilities that could lead to complete system compromise, data breach, '
        'and significant business impact.'
    )
    
    doc.add_heading('Key Findings', 2)
    
    # Findings table
    findings_table = doc.add_table(rows=6, cols=4)
    findings_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = findings_table.rows[0].cells
    header_cells[0].text = 'Finding'
    header_cells[1].text = 'Severity'
    header_cells[2].text = 'CVSS Score'
    header_cells[3].text = 'Status'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    findings_data = [
        ('Hardcoded Database Credentials', '🔴 Critical', '9.8', 'Unpatched'),
        ('SQL Injection Vulnerabilities', '🔴 Critical', '9.1', 'Unpatched'),
        ('Administrative Privilege Abuse', '🟠 High', '7.8', 'Unpatched'),
        ('No Input Validation', '🟠 High', '7.5', 'Unpatched'),
        ('Debug Mode in Production', '🟡 Medium', '6.5', 'Unpatched')
    ]
    
    for i, (finding, severity, cvss, status) in enumerate(findings_data, 1):
        row_cells = findings_table.rows[i].cells
        row_cells[0].text = finding
        row_cells[1].text = severity
        row_cells[2].text = cvss
        row_cells[3].text = status
    
    doc.add_paragraph()
    
    doc.add_heading('Risk Summary', 2)
    risk_points = [
        'Immediate Risk: Complete database compromise possible within hours',
        'Financial Impact: Estimated $4.45M average breach cost',
        'Compliance Impact: GDPR, SOC 2, PCI-DSS violations',
        'Business Impact: Operational disruption, reputational damage'
    ]
    
    for point in risk_points:
        p = doc.add_paragraph(point, style='List Bullet')
    
    doc.add_page_break()
    
    # Vulnerability Analysis
    doc.add_heading('Vulnerability Analysis', 1)
    
    # Vulnerability 1: Hardcoded Credentials
    doc.add_heading('1. Hardcoded Database Credentials (CRITICAL)', 2)
    
    doc.add_paragraph('Location: db_config.py, Lines 10-15')
    doc.add_paragraph().add_run('Vulnerable Code:').bold = True
    
    code_para = doc.add_paragraph(
        'DB_HOST = "localhost"\n'
        'DB_PORT = 5432\n'
        'DB_NAME = "inventory_db"\n'
        'DB_USER = "admin"\n'
        'DB_PASSWORD = "SuperSecret123!"  # NEVER hardcode passwords!',
        style='Intense Quote'
    )
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(9)
    
    doc.add_paragraph().add_run('Vulnerability Description:').bold = True
    
    vuln_desc = [
        'Source Code Exposure: Any developer, contractor, or attacker with repository access can view credentials',
        'Version Control History: Credentials remain in Git history even after removal',
        'No Rotation Capability: Changing passwords requires code modifications and redeployment',
        'Accidental Disclosure: Risk of exposure through code sharing, screenshots, or public repositories'
    ]
    
    for desc in vuln_desc:
        doc.add_paragraph(desc, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('CVSS v3.1 Score: 9.8 (Critical)').runs[0].font.bold = True
    
    # Vulnerability 2: SQL Injection
    doc.add_heading('2. SQL Injection Vulnerabilities (CRITICAL)', 2)
    
    doc.add_paragraph('Location: app.py, Multiple endpoints')
    doc.add_paragraph().add_run('Vulnerable Code Examples:').bold = True
    
    sql_code = doc.add_paragraph(
        '# Line 26 - Add Inventory\n'
        'cursor.execute(f"INSERT INTO inventory (item_name, quantity, warehouse_id) '
        'VALUES (\'{item_name}\', {quantity}, {warehouse_id})")\n\n'
        '# Line 88 - List Inventory\n'
        'cursor.execute(f"SELECT * FROM inventory WHERE warehouse_id = {warehouse_id}")',
        style='Intense Quote'
    )
    sql_code.runs[0].font.name = 'Courier New'
    sql_code.runs[0].font.size = Pt(9)
    
    doc.add_paragraph().add_run('Vulnerability Description:').bold = True
    doc.add_paragraph(
        'All database queries use string formatting (f-strings) to construct SQL statements, '
        'making them vulnerable to SQL injection attacks. Attackers can inject malicious SQL '
        'code through user-supplied input.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('CVSS v3.1 Score: 9.1 (Critical)').runs[0].font.bold = True
    
    doc.add_page_break()
    
    # Proof of Concept Demonstrations
    doc.add_heading('Proof of Concept Demonstrations', 1)
    
    # POC 1
    doc.add_heading('POC #1: Hardcoded Credentials Extraction', 2)
    
    doc.add_paragraph().add_run('Objective: ').bold = True
    doc.add_paragraph('Demonstrate how easily credentials can be extracted from source code')
    
    doc.add_paragraph().add_run('Attack Steps:').bold = True
    
    attack_steps = [
        'Obtain Source Code Access (clone repository, access shared code)',
        'Extract Credentials using simple Python script or regex',
        'Verify Credentials by connecting to database',
        'Confirm administrative access'
    ]
    
    for step in attack_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph().add_run('Expected Result:').bold = True
    result_para = doc.add_paragraph(
        '✓ Successfully extracted credentials:\n'
        '  Username: admin\n'
        '  Password: SuperSecret123!\n'
        '  Host: localhost\n'
        '  Database: inventory_db\n\n'
        '✓ Credentials verified - Full database access confirmed',
        style='Intense Quote'
    )
    result_para.runs[0].font.name = 'Courier New'
    result_para.runs[0].font.size = Pt(9)
    
    doc.add_paragraph().add_run('Impact: ').bold = True
    impact_para = doc.add_paragraph('Complete database compromise in under 5 minutes')
    impact_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    impact_para.runs[0].font.bold = True
    
    # POC 2
    doc.add_heading('POC #2: SQL Injection - Authentication Bypass', 2)
    
    doc.add_paragraph().add_run('Objective: ').bold = True
    doc.add_paragraph('Bypass warehouse filtering to access all inventory data')
    
    doc.add_paragraph().add_run('Vulnerable Endpoint: ').bold = True
    doc.add_paragraph('GET /inventory/list?warehouse_id={value}')
    
    doc.add_paragraph().add_run('Attack Payload:').bold = True
    payload_para = doc.add_paragraph(
        'GET /inventory/list?warehouse_id=1 OR 1=1',
        style='Intense Quote'
    )
    payload_para.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph().add_run('Resulting SQL Query:').bold = True
    sql_para = doc.add_paragraph(
        'SELECT * FROM inventory WHERE warehouse_id = 1 OR 1=1',
        style='Intense Quote'
    )
    sql_para.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph().add_run('Explanation:').bold = True
    doc.add_paragraph(
        'The "OR 1=1" condition is always true, causing the query to return ALL inventory items '
        'regardless of warehouse, effectively bypassing the intended access control.'
    )
    
    doc.add_paragraph().add_run('Impact: ').bold = True
    impact_para = doc.add_paragraph('Unauthorized access to all warehouse data')
    impact_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    impact_para.runs[0].font.bold = True
    
    doc.add_page_break()
    
    # Impact Assessment
    doc.add_heading('Impact Assessment', 1)
    
    doc.add_heading('Financial Impact', 2)
    
    financial_table = doc.add_table(rows=5, cols=2)
    financial_table.style = 'Light Grid Accent 1'
    
    financial_data = [
        ('Average Data Breach Cost', '$4.45 million (IBM 2023)'),
        ('Incident Response & Forensics', '$500,000 - $2,000,000'),
        ('GDPR Fines', 'Up to €20 million or 4% of revenue'),
        ('PCI-DSS Penalties', '$5,000 - $100,000 per month'),
        ('Ransom Payments', '$100,000 - $5,000,000')
    ]
    
    for i, (category, amount) in enumerate(financial_data):
        financial_table.rows[i].cells[0].text = category
        financial_table.rows[i].cells[1].text = amount
        financial_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('Business Impact', 2)
    
    business_impacts = [
        'System Downtime: 2-4 weeks for incident response, 4-8 weeks for rebuild',
        'Customer Trust: Immediate loss of confidence, 15-30% customer churn',
        'Reputational Damage: Negative media coverage, long-term brand impact',
        'Operational Disruption: Manual processes, order delays, supply chain issues',
        'Competitive Disadvantage: Market share loss, difficulty winning new contracts'
    ]
    
    for impact in business_impacts:
        doc.add_paragraph(impact, style='List Bullet')
    
    doc.add_page_break()
    
    # Remediation Recommendations
    doc.add_heading('Remediation Recommendations', 1)
    
    doc.add_heading('Immediate Actions (Within 24 Hours)', 2)
    
    immediate_actions = [
        'Rotate all database credentials immediately',
        'Audit database access logs for unauthorized activity',
        'Check for backdoor accounts in database',
        'Remove hardcoded credentials from source code',
        'Deploy emergency security patch',
        'Notify security team and stakeholders',
        'Document incident and actions taken'
    ]
    
    for action in immediate_actions:
        p = doc.add_paragraph(action, style='List Bullet')
        p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('Short-Term Solutions (Within 1 Week)', 2)
    
    short_term = [
        'Implement parameterized queries to prevent SQL injection',
        'Add comprehensive input validation to all endpoints',
        'Implement environment variables for configuration',
        'Update .gitignore to prevent credential commits',
        'Clean Git history to remove exposed credentials',
        'Create limited privilege database user',
        'Disable debug mode in production',
        'Test all security changes thoroughly'
    ]
    
    for action in short_term:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Long-Term Solutions (Within 3 Months)', 2)
    
    long_term = [
        'Deploy secrets management system (HashiCorp Vault, AWS Secrets Manager)',
        'Implement automatic credential rotation',
        'Enable database encryption (TLS/SSL)',
        'Deploy SIEM solution for security monitoring',
        'Implement automated security scanning in CI/CD',
        'Conduct penetration testing',
        'Implement comprehensive security training program',
        'Establish incident response procedures'
    ]
    
    for action in long_term:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_page_break()
    
    # Secure Code Example
    doc.add_heading('Secure Code Implementation Example', 1)
    
    doc.add_paragraph().add_run('Secure Endpoint Implementation:').bold = True
    
    secure_code = doc.add_paragraph(
        '@app.route(\'/inventory/add\', methods=[\'POST\'])\n'
        '@limiter.limit("10 per minute")\n'
        '@token_required\n'
        'def add_inventory(current_user):\n'
        '    schema = InventorySchema()\n'
        '    \n'
        '    try:\n'
        '        # Validate input\n'
        '        data = schema.load(request.get_json())\n'
        '    except ValidationError as err:\n'
        '        return jsonify({\'errors\': err.messages}), 400\n'
        '    \n'
        '    conn = get_db_connection()\n'
        '    cursor = conn.cursor()\n'
        '    \n'
        '    # Parameterized query - SECURE\n'
        '    cursor.execute(\n'
        '        "INSERT INTO inventory (item_name, quantity, warehouse_id) VALUES (?, ?, ?)",\n'
        '        (data[\'item_name\'], data[\'quantity\'], data[\'warehouse_id\'])\n'
        '    )\n'
        '    \n'
        '    conn.commit()\n'
        '    return jsonify({\'status\': \'success\'}), 201',
        style='Intense Quote'
    )
    secure_code.runs[0].font.name = 'Courier New'
    secure_code.runs[0].font.size = Pt(8)
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    
    conclusion_text = (
        'The credential management vulnerabilities in db_config.py and SQL injection vulnerabilities '
        'in app.py represent CRITICAL security risks that require immediate remediation. The combination '
        'of hardcoded administrative credentials, plaintext storage, and lack of security controls creates '
        'multiple attack vectors that could lead to:\n\n'
        '• Complete database compromise\n'
        '• Significant financial losses ($4M+ average)\n'
        '• Severe reputational damage\n'
        '• Regulatory penalties and legal liability\n'
        '• Business continuity disruption\n\n'
        'IMMEDIATE ACTION IS REQUIRED to rotate credentials, implement parameterized queries, '
        'and deploy comprehensive security controls. Failure to address these vulnerabilities promptly '
        'could result in a catastrophic security incident with severe consequences for the organization.'
    )
    
    conclusion_para = doc.add_paragraph(conclusion_text)
    conclusion_para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Warning box
    warning = doc.add_paragraph()
    warning_run = warning.add_run(
        '⚠️ CRITICAL WARNING: This system should NOT be deployed to production until all '
        'identified vulnerabilities are remediated and verified through security testing.'
    )
    warning_run.font.bold = True
    warning_run.font.color.rgb = RGBColor(255, 0, 0)
    warning_run.font.size = Pt(12)
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Sign-off
    doc.add_heading('Document Approval and Sign-off', 1)
    
    signoff_table = doc.add_table(rows=4, cols=2)
    signoff_table.style = 'Light Grid Accent 1'
    
    signoff_data = [
        ('Prepared By:', 'Bob (AI Security Analyst)'),
        ('Date:', 'February 5, 2026'),
        ('Classification:', 'CONFIDENTIAL - Internal Use Only'),
        ('Document Version:', '1.0')
    ]
    
    for i, (key, value) in enumerate(signoff_data):
        signoff_table.rows[i].cells[0].text = key
        signoff_table.rows[i].cells[1].text = value
        signoff_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Save document
    output_path = 'Security_Vulnerability_Analysis_Report.docx'
    doc.save(output_path)
    print(f"✓ Word document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_security_report()

# Made with Bob
